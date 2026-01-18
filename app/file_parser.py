"""
File content parser for different file types (PDF, DOCX, TXT, code files)
"""
import os
from typing import Optional
from pathlib import Path

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text content from various file types.

    Args:
        file_path: Path to the file

    Returns:
        Extracted text content
    """
    if not os.path.exists(file_path):
        return "[File not found]"

    file_ext = Path(file_path).suffix.lower()

    try:
        # PDF files
        if file_ext == '.pdf':
            return _extract_from_pdf(file_path)

        # DOCX files
        elif file_ext in ['.docx', '.doc']:
            return _extract_from_docx(file_path)

        # Text-based files (code, txt, md, etc.)
        elif file_ext in ['.txt', '.py', '.java', '.cpp', '.c', '.js', '.html', '.css',
                          '.json', '.xml', '.md', '.csv', '.sh', '.sql', '.r', '.m',
                          '.swift', '.kt', '.rs', '.go', '.rb', '.php', '.yaml', '.yml']:
            return _extract_from_text(file_path)

        # Try as text file if unknown extension
        else:
            try:
                return _extract_from_text(file_path)
            except Exception:
                return f"[Unsupported file type: {file_ext}. File size: {os.path.getsize(file_path)} bytes]"

    except Exception as e:
        return f"[Error reading file: {str(e)}]"


def _extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using PyPDF2"""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        text_parts = []

        # Extract text from each page
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text.strip():
                text_parts.append(f"--- Page {i+1} ---\n{page_text}")

        if not text_parts:
            return "[PDF contains no extractable text. It may be an image-based PDF or encrypted.]"

        full_text = "\n\n".join(text_parts)

        # Add metadata
        metadata = []
        if reader.metadata:
            if reader.metadata.title:
                metadata.append(f"Title: {reader.metadata.title}")
            if reader.metadata.author:
                metadata.append(f"Author: {reader.metadata.author}")

        metadata.append(f"Total Pages: {len(reader.pages)}")

        if metadata:
            return f"PDF Metadata:\n{chr(10).join(metadata)}\n\n{full_text}"

        return full_text

    except Exception as e:
        return f"[Error parsing PDF: {str(e)}]"


def _extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX file using python-docx"""
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        if not text_parts:
            return "[DOCX contains no extractable text]"

        return "\n\n".join(text_parts)

    except Exception as e:
        return f"[Error parsing DOCX: {str(e)}]"


def _extract_from_text(file_path: str) -> str:
    """Extract text from plain text files"""
    # Try different encodings
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue

    # If all encodings fail, read as binary and decode with errors='ignore'
    try:
        with open(file_path, 'rb') as f:
            return f.read().decode('utf-8', errors='ignore')
    except Exception as e:
        return f"[Error reading text file: {str(e)}]"


def get_file_info(file_path: str) -> dict:
    """Get file information"""
    if not os.path.exists(file_path):
        return {"exists": False}

    stat = os.stat(file_path)
    return {
        "exists": True,
        "size_bytes": stat.st_size,
        "size_kb": round(stat.st_size / 1024, 2),
        "extension": Path(file_path).suffix.lower(),
        "name": Path(file_path).name
    }

